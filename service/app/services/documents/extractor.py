from __future__ import annotations

import re
from dataclasses import dataclass
from typing import Sequence

import pymupdf
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Basic lookup tables to derive domain + seniority from the source CV text.
EMAIL_RE = re.compile(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", re.IGNORECASE)
PHONE_RE = re.compile(r"\b(?:\+?\d{1,3}[\s.-]?)?(?:\(?\d{3}\)?[\s.-]?){2}\d{4}\b")
LINK_RE = re.compile(r"\bhttps?://[\w\-.]*linkedin\.com/[^\s]+", re.IGNORECASE)
# Street address pattern (e.g., "123 Main Street")
ADDRESS_RE = re.compile(r"\b\d{2,5}\s+[A-Z0-9][\w\s.-]+(?:Street|St|Road|Rd|Avenue|Ave|Boulevard|Blvd|Lane|Ln|Drive|Dr|Court|Ct|Place|Pl)\b", re.IGNORECASE)
# City, State/Province, Postal Code pattern (e.g., "Doncaster, Victoria 3106" or "New York, NY 10001")
# Matches: City name, comma, state/province (2+ chars), optional postal code (4-5 digits)
# This is intentionally broad to catch various location formats
LOCATION_RE = re.compile(
    r"\b[A-Z][a-zA-Z\s]{2,25},\s*[A-Z][a-zA-Z\s]{2,20}(?:\s+\d{4,5}(?:-\d{4})?)?\b",
    re.IGNORECASE
)
# GitHub URLs
GITHUB_RE = re.compile(r"\b(?:https?://)?(?:www\.)?github\.com/[\w\-]+(?:\s|$)", re.IGNORECASE)
# General URLs (excluding already matched LinkedIn and GitHub)
URL_RE = re.compile(r"\bhttps?://(?![\w\-.]*linkedin\.com|[\w\-.]*github\.com)[^\s]+", re.IGNORECASE)
# SSN pattern (XXX-XX-XXXX or XXX XX XXXX)
SSN_RE = re.compile(r"\b\d{3}[-.\s]?\d{2}[-.\s]?\d{4}\b")
# Passport pattern (alphanumeric, typically 6-9 characters)
PASSPORT_RE = re.compile(r"\b[A-Z]{1,2}\d{6,9}\b")
# Date of birth patterns (MM/DD/YYYY, DD/MM/YYYY, YYYY-MM-DD, etc.)
DOB_RE = re.compile(r"\b(?:0?[1-9]|1[0-2])[/-](?:0?[1-9]|[12]\d|3[01])[/-](?:19|20)\d{2}\b")
# Credit card pattern (basic - 16 digits with optional separators)
CC_RE = re.compile(r"\b\d{4}[-.\s]?\d{4}[-.\s]?\d{4}[-.\s]?\d{4}\b")

REDACTION_MAP = {
    EMAIL_RE: "[redacted-email]",
    PHONE_RE: "[redacted-phone]",
    LINK_RE: "[redacted-link]",
    GITHUB_RE: "[redacted-link]",
    URL_RE: "[redacted-link]",
    ADDRESS_RE: "[redacted-address]",
    LOCATION_RE: "[redacted-location]",
    SSN_RE: "[redacted-ssn]",
    PASSPORT_RE: "[redacted-passport]",
    DOB_RE: "[redacted-dob]",
    CC_RE: "[redacted-cc]",
}


@dataclass(slots=True)
class CreatorStudioExtraction:
    name: str
    sanitized_text: str
    confidence_score: float = 1.0
    redaction_count: int = 0

    def to_payload(self) -> dict[str, str | float | int]:
        return {
            "name": self.name,
            "sanitized_text": self.sanitized_text,
            "confidence_score": self.confidence_score,
            "redaction_count": self.redaction_count,
        }


class CreatorStudioExtractor:
    """Parse contributor PDFs/DOCX files and output PII-sanitised Knight fields."""

    def parse(self, file_bytes: bytes, filename: str = "upload.pdf") -> CreatorStudioExtraction:
        text = self._extract_text(file_bytes, filename)
        if not text.strip():
            raise ValueError("The uploaded file did not contain extractable text.")
        sanitized, redaction_count = self._redact_pii(text)
        lines = [line.strip() for line in sanitized.splitlines() if line.strip()]

        name = self._guess_name(lines)
        
        # Redact the extracted name from sanitized text (if found)
        # This ensures names are removed from the text sent to LLM
        if name and name != "Resume-sourced Knight":
            # Escape special regex characters in the name
            name_pattern = re.escape(name)
            # Match the name as a whole word or at start/end of line
            name_re = re.compile(rf"\b{name_pattern}\b", re.IGNORECASE)
            matches = len(name_re.findall(sanitized))
            sanitized = name_re.sub("[redacted-name]", sanitized)
            redaction_count += matches
        
        sanitized_text = self._build_sanitized_text(sanitized)
        
        # Calculate confidence score based on text quality and redactions
        confidence = self._calculate_confidence(text, sanitized_text, redaction_count)

        return CreatorStudioExtraction(
            name=name,
            sanitized_text=sanitized_text,
            confidence_score=confidence,
            redaction_count=redaction_count,
        )

    def _extract_text(self, file_bytes: bytes, filename: str) -> str:
        filename_lower = filename.lower()
        if filename_lower.endswith(".docx"):
            if not DOCX_AVAILABLE:
                raise ValueError("DOCX support requires python-docx library. Please install it.")
            try:
                from io import BytesIO
                doc = Document(BytesIO(file_bytes))
                parts: list[str] = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        parts.append(paragraph.text)
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_text:
                            parts.append(row_text)
                return "\n".join(parts)
            except Exception as exc:
                raise ValueError("Unable to open the DOCX payload.") from exc
        else:
            # Default to PDF
            try:
                with pymupdf.open(stream=file_bytes, filetype="pdf") as document:
                    parts: list[str] = []
                    for page in document:
                        page_text = page.get_text("text")
                        if page_text:
                            parts.append(page_text)
            except Exception as exc:  # pragma: no cover - low-level errors depend on MuPDF build
                raise ValueError("Unable to open the PDF payload.") from exc
            return "\n".join(parts)

    def _redact_pii(self, text: str) -> tuple[str, int]:
        redacted = text
        total_redactions = 0
        for pattern, replacement in REDACTION_MAP.items():
            matches = len(pattern.findall(redacted))
            redacted = pattern.sub(replacement, redacted)
            total_redactions += matches
        return redacted, total_redactions

    def _calculate_confidence(self, original: str, sanitized: str, redaction_count: int) -> float:
        """Calculate confidence score (0.0 to 1.0) based on extraction quality."""
        if not original or not sanitized:
            return 0.0
        
        # Base confidence from text length (more text = higher confidence)
        text_length_score = min(len(sanitized) / 500.0, 1.0)  # Normalize to 500 chars
        
        # Penalty for excessive redactions (might indicate poor extraction)
        redaction_penalty = min(redaction_count / 10.0, 0.3)  # Max 30% penalty
        
        # Check if we have meaningful content (not just redactions)
        meaningful_content = len(sanitized.replace("[redacted-", "")) / len(sanitized) if sanitized else 0
        
        confidence = (text_length_score * 0.4 + meaningful_content * 0.6) - redaction_penalty
        return max(0.0, min(1.0, confidence))

    def _guess_name(self, lines: Sequence[str]) -> str:
        for entry in lines[:5]:
            if self._looks_like_name(entry):
                return entry
        return "Resume-sourced Knight"

    @staticmethod
    def _looks_like_name(value: str) -> bool:
        tokens = value.split()
        if len(tokens) not in (2, 3):
            return False
        if any(char.isdigit() for char in value):
            return False
        if "@" in value.lower():
            return False
        uppercase_tokens = sum(1 for token in tokens if token.isupper())
        # Allow typical "John Doe" or "JANE DOE" styles.
        return uppercase_tokens <= len(tokens)

    def _build_sanitized_text(self, text: str) -> str:
        return text.strip()
